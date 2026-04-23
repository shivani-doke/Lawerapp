# backend/routes/document_routes.py

import os
import json
import tempfile
import uuid
import re
import shutil
from html import escape, unescape
from html.parser import HTMLParser
from datetime import datetime
from flask import Blueprint, request, jsonify, send_from_directory, abort
from werkzeug.utils import secure_filename

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt
import PyPDF2

from services.gemini_service import (
    SUPPORTED_GENERATION_LANGUAGES,
    fill_document_with_fields,
    generate_document_from_fields_only,
    localize_field_schema,
    localize_field_values,
    replace_placeholders,
)
from services.document_service import (
    GENERATED_DOCS_FOLDER,
    GENERATED_FOLDER,
    PREVIEW_FOLDER,
    convert_docx_to_pdf,
)
from services.document_fields import get_fields_for_document_type, get_subtypes_for_document_type
from services.auth_context import get_current_user, get_request_username
from models.app_user_model import AppUser

SUPPORTED_FONT_FAMILIES = (
    "Times New Roman",
    "Arial",
    "Calibri",
    "Cambria",
    "Book Antiqua",
    "Bookman Old Style",
    "Candara",
    "Century Gothic",
    "Comic Sans MS",
    "Consolas",
    "Constantia",
    "Corbel",
    "Courier New",
    "Georgia",
    "Garamond",
    "Lucida Console",
    "Lucida Sans Unicode",
    "Palatino Linotype",
    "Segoe UI",
    "Sylfaen",
    "Trebuchet MS",
    "Verdana",
    "Tahoma",
    "Nirmala UI",
    "Mangal",
    "Aparajita",
    "Kokila",
    "Utsaah",
)

DEFAULT_FONT_SIZE = 14
MIN_FONT_SIZE = 8
MAX_FONT_SIZE = 72

# Folder where template DOCX files are stored
TEMPLATES_FOLDER = "templates"
os.makedirs(TEMPLATES_FOLDER, exist_ok=True)

# Folder for uploaded reference documents
UPLOADS_FOLDER = "uploads"
os.makedirs(UPLOADS_FOLDER, exist_ok=True)
METADATA_FILE = os.path.join(UPLOADS_FOLDER, "metadata.json")
GENERATED_METADATA_FILE = os.path.join(GENERATED_FOLDER, "generated_metadata.json")
SMARTLEGAL_FOLDER = os.path.join(UPLOADS_FOLDER, "smartlegal")
os.makedirs(SMARTLEGAL_FOLDER, exist_ok=True)
SMARTLEGAL_METADATA_FILE = os.path.join(SMARTLEGAL_FOLDER, "metadata.json")

document_bp = Blueprint("document_bp", __name__)

# -------------------------------------------------------------------
# Helper functions for metadata
# -------------------------------------------------------------------
def load_metadata():
    if os.path.exists(METADATA_FILE):
        with open(METADATA_FILE, 'r') as f:
            return json.load(f)
    return {}

def save_metadata(metadata):
    with open(METADATA_FILE, 'w') as f:
        json.dump(metadata, f, indent=2)


def load_generated_metadata():
    if os.path.exists(GENERATED_METADATA_FILE):
        with open(GENERATED_METADATA_FILE, "r") as f:
            return json.load(f)
    return {}


def save_generated_metadata(metadata):
    with open(GENERATED_METADATA_FILE, "w") as f:
        json.dump(metadata, f, indent=2)


def load_smartlegal_metadata():
    if os.path.exists(SMARTLEGAL_METADATA_FILE):
        with open(SMARTLEGAL_METADATA_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}


def save_smartlegal_metadata(metadata):
    with open(SMARTLEGAL_METADATA_FILE, "w", encoding="utf-8") as f:
        json.dump(metadata, f, indent=2)


def resolve_firm_name(owner_username):
    if not owner_username:
        return "Default Firm"
    user = AppUser.query.filter_by(username=owner_username).first()
    if user and (user.firm_name or "").strip():
        return user.firm_name.strip()
    return "Default Firm"


def resolve_firm_id(owner_username):
    if not owner_username:
        return None
    user = AppUser.query.filter_by(username=owner_username).first()
    return user.firm_id if user else None


def metadata_firm_name(info):
    info = info or {}
    return (
        info.get("firm_name")
        or resolve_firm_name(info.get("owner_username"))
        or "Default Firm"
    )


def metadata_firm_id(info):
    info = info or {}
    if info.get("firm_id") is not None:
        return info.get("firm_id")
    return resolve_firm_id(info.get("owner_username"))


def metadata_accessible_to_user(info, user):
    if user is None:
        return False
    if metadata_firm_id(info) is not None and user.firm_id is not None:
        return metadata_firm_id(info) == user.firm_id
    return metadata_firm_name(info) == user.firm_name


def ensure_generated_metadata_defaults():
    metadata = load_generated_metadata()
    changed = False

    for folder in (GENERATED_FOLDER, GENERATED_DOCS_FOLDER):
        if not os.path.exists(folder):
            continue
        for filename in os.listdir(folder):
            full_path = os.path.join(folder, filename)
            if not os.path.isfile(full_path):
                continue
            if filename not in metadata:
                metadata[filename] = {
                    "owner_username": "admin",
                    "firm_id": 1,
                    "firm_name": "Default Firm",
                    "timestamp": datetime.fromtimestamp(
                        os.path.getctime(full_path)
                    ).isoformat(),
                }
                changed = True
            else:
                if not metadata[filename].get("owner_username"):
                    metadata[filename]["owner_username"] = "admin"
                    changed = True
                if not metadata[filename].get("firm_name"):
                    metadata[filename]["firm_name"] = resolve_firm_name(
                        metadata[filename].get("owner_username")
                    )
                    changed = True
                if metadata[filename].get("firm_id") is None:
                    metadata[filename]["firm_id"] = resolve_firm_id(
                        metadata[filename].get("owner_username")
                    )
                    changed = True

    if changed:
        save_generated_metadata(metadata)

def normalize_document_type(document_type):
    return (document_type or "").strip().lower()

def get_document_title(document_type):
    doc_type = normalize_document_type(document_type)
    title_map = {
        "power_of_attorney": "POWER OF ATTORNEY",
        "gift_deed": "GIFT DEED",
        "rental_agreement": "RENTAL AGREEMENT",
        "partnership_deed": "PARTNERSHIP DEED",
        "affidavit": "AFFIDAVIT",
        "will_and_testament": "WILL AND TESTAMENT",
        "bail_application": "BAIL APPLICATION",
        "loan_agreement": "LOAN AGREEMENT",
        "divorce_paper": "DIVORCE PAPER",
        "sale_deed": "SALE DEED",
        "mortgage_deed": "MORTGAGE DEED",
        "non_disclosure_agreement": "NON-DISCLOSURE AGREEMENT",
        "employment_contract": "EMPLOYMENT CONTRACT",
        "offer_letter": "OFFER LETTER",
        "service_agreement": "SERVICE AGREEMENT",
        "child_custody_agreement": "CHILD CUSTODY AGREEMENT",
        "adoption_papers": "ADOPTION PAPERS",
        "partition_deed": "PARTITION DEED",
        "trust_deed": "TRUST DEED",
        "memorandum_of_understanding": "MEMORANDUM OF UNDERSTANDING",
        "vendor_agreement": "VENDOR AGREEMENT",
        "non_compete_agreement": "NON-COMPETE AGREEMENT",
        "indemnity_agreement": "INDEMNITY AGREEMENT",
        "joint_venture_agreement": "JOINT VENTURE AGREEMENT",
        "licensing_agreement": "LICENSING AGREEMENT",
        "assignment_agreement": "ASSIGNMENT AGREEMENT",
        "settlement_agreement": "SETTLEMENT AGREEMENT",
        "trademark_application": "TRADEMARK APPLICATION",
        "copyright_agreement": "COPYRIGHT AGREEMENT",
        "patent_filing_documents": "PATENT FILING DOCUMENTS",
    }
    if doc_type in title_map:
        return title_map[doc_type]
    return doc_type.replace("_", " ").strip().upper()


def get_generated_file_directory(filename):
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        return GENERATED_FOLDER
    if ext == ".docx":
        return GENERATED_DOCS_FOLDER
    return GENERATED_FOLDER


def normalize_generation_language(language):
    requested = (language or "").strip().lower()
    for supported_language in SUPPORTED_GENERATION_LANGUAGES:
        if requested == supported_language.lower():
            return supported_language
    return SUPPORTED_GENERATION_LANGUAGES[0]


def normalize_font_family(font_family):
    requested = (font_family or "").strip().lower()
    for supported_font_family in SUPPORTED_FONT_FAMILIES:
        if requested == supported_font_family.lower():
            return supported_font_family
    return SUPPORTED_FONT_FAMILIES[0]


def normalize_font_size(font_size):
    try:
        parsed_size = int(str(font_size).strip())
    except (TypeError, ValueError):
        return DEFAULT_FONT_SIZE
    return max(MIN_FONT_SIZE, min(MAX_FONT_SIZE, parsed_size))


def normalize_paper_size(paper_size):
    requested = str(paper_size or "").strip().lower()
    supported_sizes = {
        "a4": "A4",
        "letter": "Letter",
        "legal": "Legal",
    }
    return supported_sizes.get(requested, "A4")


def normalize_line_spacing(line_spacing):
    requested = str(line_spacing or "").strip().lower()
    supported_spacing = {
        "single": 1.0,
        "1": 1.0,
        "1.0": 1.0,
        "1.15": 1.15,
        "1.5": 1.5,
        "double": 2.0,
        "2": 2.0,
        "2.0": 2.0,
    }
    return supported_spacing.get(requested, 1.0)


def line_spacing_label(line_spacing):
    spacing = normalize_line_spacing(line_spacing)
    if spacing == 1.0:
        return "Single"
    if spacing == 1.15:
        return "1.15"
    if spacing == 1.5:
        return "1.5"
    if spacing == 2.0:
        return "Double"
    return "Single"


def normalize_margin_size(margin_size):
    requested = str(margin_size or "").strip().lower()
    supported_margins = {
        "normal": "Normal",
        "narrow": "Narrow",
        "moderate": "Moderate",
        "wide": "Wide",
    }
    return supported_margins.get(requested, "Normal")


def apply_font_to_run(run, font_family, font_size):
    run.font.name = font_family
    run.font.size = Pt(font_size)
    run_properties = run._element.get_or_add_rPr()
    run_properties.rFonts.set(qn("w:ascii"), font_family)
    run_properties.rFonts.set(qn("w:hAnsi"), font_family)
    run_properties.rFonts.set(qn("w:eastAsia"), font_family)
    run_properties.rFonts.set(qn("w:cs"), font_family)


def apply_default_font_style(doc, font_family, font_size):
    normal_style = doc.styles["Normal"]
    normal_style.font.name = font_family
    normal_style.font.size = Pt(font_size)
    style_properties = normal_style.element.get_or_add_rPr()
    style_properties.rFonts.set(qn("w:ascii"), font_family)
    style_properties.rFonts.set(qn("w:hAnsi"), font_family)
    style_properties.rFonts.set(qn("w:eastAsia"), font_family)
    style_properties.rFonts.set(qn("w:cs"), font_family)


def apply_font_settings(doc, font_family, font_size):
    apply_default_font_style(doc, font_family, font_size)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            apply_font_to_run(run, font_family, font_size)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        apply_font_to_run(run, font_family, font_size)


def apply_paper_size(doc, paper_size):
    for section in doc.sections:
        if paper_size == "Letter":
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
        elif paper_size == "Legal":
            section.page_width = Inches(8.5)
            section.page_height = Inches(14)
        elif paper_size == "A4":
            section.page_width = Mm(210)
            section.page_height = Mm(297)
        else:
            # Handle custom sizes like "210.0x297.0mm"
            import re
            match = re.match(r'(\d+(?:\.\d+)?)x(\d+(?:\.\d+)?)mm', paper_size)
            if match:
                width_mm = float(match.group(1))
                height_mm = float(match.group(2))
                section.page_width = Mm(width_mm)
                section.page_height = Mm(height_mm)
            else:
                # Fallback to A4
                section.page_width = Mm(210)
                section.page_height = Mm(297)


def apply_margin_settings(doc, margin_size):
    margin_map = {
        "Normal": Inches(1.0),
        "Narrow": Inches(0.5),
        "Moderate": Inches(0.75),
        "Wide": Inches(1.5),
    }
    margin_value = margin_map.get(margin_size, margin_map["Normal"])

    for section in doc.sections:
        section.top_margin = margin_value
        section.bottom_margin = margin_value
        section.left_margin = margin_value
        section.right_margin = margin_value


def apply_line_spacing_settings(doc, line_spacing):
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing = line_spacing

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = line_spacing


def apply_document_layout_settings(doc, paper_size, line_spacing, margin_size, apply_line_spacing=True):
    apply_paper_size(doc, paper_size)
    apply_margin_settings(doc, margin_size)
    if apply_line_spacing:
        apply_line_spacing_settings(doc, line_spacing)

def enrich_divorce_fields(fields):
    enriched = dict(fields or {})
    divorce_type = str(enriched.get("divorce_type", "")).strip()
    mutual_first_petitioner = str(enriched.get("mutual_first_petitioner", "")).strip()
    contested_filed_by = str(enriched.get("contested_filed_by", "")).strip()

    husband_name = str(enriched.get("husband_name", "")).strip()
    wife_name = str(enriched.get("wife_name", "")).strip()

    if divorce_type == "Mutual Consent":
        first_party = mutual_first_petitioner or "Husband"
        second_party = "Wife" if first_party == "Husband" else "Husband"

        enriched["husband_role_label"] = (
            "Petitioner No. 1" if first_party == "Husband" else "Petitioner No. 2"
        )
        enriched["wife_role_label"] = (
            "Petitioner No. 1" if first_party == "Wife" else "Petitioner No. 2"
        )
        enriched["petitioner_no_1_party"] = first_party
        enriched["petitioner_no_2_party"] = second_party
        enriched["petitioner_no_1_name"] = husband_name if first_party == "Husband" else wife_name
        enriched["petitioner_no_2_name"] = wife_name if second_party == "Wife" else husband_name
        enriched["party_role_summary"] = (
            f"Mutual consent divorce. {first_party} is Petitioner No. 1 and "
            f"{second_party} is Petitioner No. 2."
        )
        enriched["petitioner_name"] = enriched["petitioner_no_1_name"]
        enriched["respondent_name"] = enriched["petitioner_no_2_name"]

    elif divorce_type == "Contested":
        filing_party = contested_filed_by or "Husband"
        responding_party = "Wife" if filing_party == "Husband" else "Husband"

        enriched["husband_role_label"] = (
            "Petitioner" if filing_party == "Husband" else "Respondent"
        )
        enriched["wife_role_label"] = (
            "Petitioner" if filing_party == "Wife" else "Respondent"
        )
        enriched["petitioner_party"] = filing_party
        enriched["respondent_party"] = responding_party
        enriched["petitioner_name"] = husband_name if filing_party == "Husband" else wife_name
        enriched["respondent_name"] = wife_name if responding_party == "Wife" else husband_name
        enriched["party_role_summary"] = (
            f"Contested divorce. {filing_party} is the Petitioner and "
            f"{responding_party} is the Respondent."
        )

    return enriched

def extract_text_from_file(filepath, ext):
    """Extract text from a file based on its extension."""
    try:
        if ext == ".txt":
            with open(filepath, "r", encoding="utf-8") as f:
                return f.read()
        elif ext == ".docx":
            doc = DocxDocument(filepath)
            return "\n".join([para.text for para in doc.paragraphs])
        elif ext == ".pdf":
            with open(filepath, "rb") as f:
                pdf_reader = PyPDF2.PdfReader(f)
                extracted_pages = []
                for page in pdf_reader.pages:
                    page_text = page.extract_text(Tj_sep=" ", TJ_sep=" ")
                    extracted_pages.append((page_text or "").strip("\n"))
                return "\n\n".join(extracted_pages)
        else:
            return None
    except Exception as e:
        print(f"Text extraction error: {e}")
        return None


def clear_document_body(doc):
    body = doc.element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def get_generated_html_filepath(filename):
    stem, _ = os.path.splitext(filename)
    return os.path.join(GENERATED_DOCS_FOLDER, f"{stem}.editor.html")


def save_generated_html_content(filename, html_content):
    html_path = get_generated_html_filepath(filename)
    with open(html_path, "w", encoding="utf-8") as html_file:
        html_file.write(html_content or "")


def load_generated_html_content(filename):
    html_path = get_generated_html_filepath(filename)
    if not os.path.exists(html_path):
        return None
    with open(html_path, "r", encoding="utf-8") as html_file:
        return html_file.read()


def parse_serialized_fields(raw_value, fallback_fields):
    if raw_value is None:
        return fallback_fields
    try:
        parsed = json.loads(raw_value)
    except Exception:
        return fallback_fields
    return parsed if isinstance(parsed, list) else fallback_fields


def get_source_reference_docx_path(info):
    source_filename = str((info or {}).get("source_reference_filename") or "").strip()
    if not source_filename.lower().endswith(".docx"):
        return None
    source_path = os.path.join(UPLOADS_FOLDER, source_filename)
    if not os.path.exists(source_path):
        return None
    return source_path


def get_matching_generated_docx_filename(filename):
    stem, ext = os.path.splitext(filename)
    if ext.lower() == ".docx":
        return filename
    if ext.lower() == ".pdf":
        return f"{stem}.docx"
    return None


def build_html_from_docx(filepath):
    doc = DocxDocument(filepath)
    blocks = []

    for paragraph in doc.paragraphs:
        tag = "p"
        if paragraph.style is not None:
            style_name = str(paragraph.style.name or "").lower()
            if "heading 1" in style_name:
                tag = "h1"
            elif "heading 2" in style_name:
                tag = "h2"
            elif "heading 3" in style_name:
                tag = "h3"

        styles = []
        if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            styles.append("text-align:center")
        elif paragraph.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            styles.append("text-align:right")
        elif paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            styles.append("text-align:justify")

        if paragraph.paragraph_format.left_indent:
            indent_pt = paragraph.paragraph_format.left_indent.pt
            if indent_pt:
                styles.append(f"margin-left:{indent_pt:.0f}pt")

        segments = []
        if paragraph.runs:
            for run in paragraph.runs:
                text = escape(run.text or "")
                if not text and not run._element.xpath(".//w:br"):
                    continue
                text = text.replace("\n", "<br>")
                if run._element.xpath(".//w:br"):
                    text += "<br>"

                inline_styles = []
                if run.font.name:
                    inline_styles.append(f"font-family:{escape(run.font.name)}")
                if run.font.size:
                    inline_styles.append(f"font-size:{run.font.size.pt:.0f}px")
                if run.bold:
                    inline_styles.append("font-weight:bold")
                if run.italic:
                    inline_styles.append("font-style:italic")
                if run.underline:
                    inline_styles.append("text-decoration:underline")
                if run.font.strike:
                    inline_styles.append("text-decoration:line-through")

                if inline_styles:
                    segments.append(
                        f'<span style="{"; ".join(inline_styles)}">{text}</span>'
                    )
                else:
                    segments.append(text)
        else:
            segments.append("<br>")

        style_attr = f' style="{"; ".join(styles)}"' if styles else ""
        blocks.append(f"<{tag}{style_attr}>{''.join(segments) or '<br>'}</{tag}>")

    return "".join(blocks) if blocks else "<p></p>"


def html_to_plain_text(html_content):
    html_content = str(html_content or "")
    if not html_content.strip():
        return ""
    text = re.sub(r"(?i)<br\s*/?>", "\n", html_content)
    text = re.sub(r"(?i)</(p|div|h1|h2|h3|li|ul|ol|blockquote)>", "\n", text)
    text = re.sub(r"(?i)<hr[^>]*>", "\n----------------------------------------\n", text)
    text = re.sub(r"<[^>]+>", "", text)
    text = text.replace("&nbsp;", " ")
    text = text.replace("&amp;", "&")
    text = text.replace("&lt;", "<")
    text = text.replace("&gt;", ">")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def parse_editor_font_size(value, default_size):
    size_map = {
        "1": 8,
        "2": 10,
        "3": 12,
        "4": 14,
        "5": 18,
        "6": 24,
        "7": 32,
    }
    parsed = str(value or "").strip()
    return size_map.get(parsed, default_size)


def parse_inline_style(style_value):
    style_map = {}
    for declaration in str(style_value or "").split(";"):
        if ":" not in declaration:
            continue
        key, value = declaration.split(":", 1)
        style_map[key.strip().lower()] = value.strip()
    return style_map


class EditorHtmlParser(HTMLParser):
    def __init__(self, default_font_family, default_font_size):
        super().__init__(convert_charrefs=False)
        self.default_font_family = default_font_family
        self.default_font_size = default_font_size
        self.blocks = []
        self.current_block = None
        self.style_stack = [self._base_style()]
        self.list_stack = []
        self.indent_stack = [0.0]

    def _base_style(self):
        return {
            "bold": False,
            "italic": False,
            "underline": False,
            "strike": False,
            "font_family": self.default_font_family,
            "font_size": self.default_font_size,
        }

    def _current_style(self):
        return dict(self.style_stack[-1])

    def _extract_alignment(self, attrs):
        style_map = parse_inline_style(attrs.get("style"))
        alignment = (style_map.get("text-align") or "").strip().lower()
        if alignment in {"center", "right", "left", "justify"}:
            return alignment
        return None

    def _extract_indent_points(self, attrs):
        style_map = parse_inline_style(attrs.get("style"))
        margin_left = (style_map.get("margin-left") or "").strip().lower()
        if not margin_left:
            return 0.0

        match = re.search(r"(\d+(?:\.\d+)?)", margin_left)
        if not match:
            return 0.0

        numeric = float(match.group(1))
        if margin_left.endswith("px"):
            return numeric * 0.75
        if margin_left.endswith("pt"):
            return numeric
        if margin_left.endswith("in"):
            return numeric * 72.0
        if margin_left.endswith("cm"):
            return numeric * 28.3465
        if margin_left.endswith("mm"):
            return numeric * 2.83465
        return numeric

    def _extract_space_before_points(self, attrs):
        style_map = parse_inline_style(attrs.get("style"))
        margin_top = (
            style_map.get("margin-top")
            or style_map.get("padding-top")
            or ""
        ).strip().lower()
        if not margin_top:
            return 0.0

        match = re.search(r"(\d+(?:\.\d+)?)", margin_top)
        if not match:
            return 0.0

        numeric = float(match.group(1))
        if margin_top.endswith("px"):
            return numeric * 0.75
        if margin_top.endswith("pt"):
            return numeric
        if margin_top.endswith("in"):
            return numeric * 72.0
        if margin_top.endswith("cm"):
            return numeric * 28.3465
        if margin_top.endswith("mm"):
            return numeric * 2.83465
        return numeric

    def _extract_line_spacing(self, attrs):
        style_map = parse_inline_style(attrs.get("style"))
        line_height = (style_map.get("line-height") or "").strip().lower()
        if not line_height:
            return None

        match = re.search(r"(\d+(?:\.\d+)?)", line_height)
        if not match:
            return None

        numeric = float(match.group(1))
        if numeric <= 0:
            return None

        if any(
            line_height.endswith(unit)
            for unit in ("px", "pt", "in", "cm", "mm", "%")
        ):
            return None
        return numeric

    def _current_indent(self):
        return float(self.indent_stack[-1])

    def _start_block(self, tag, attrs):
        if self.current_block is not None:
            self._finalize_block()
        self.current_block = {
            "tag": tag,
            "align": self._extract_alignment(attrs),
            "indent": self._current_indent() + self._extract_indent_points(attrs),
            "space_before": self._extract_space_before_points(attrs),
            "line_spacing": self._extract_line_spacing(attrs),
            "segments": [],
        }

    def _append_text(self, text):
        if text is None:
            return
        if self.current_block is None:
            self._start_block("p", {})
        normalized = unescape(text)
        if not normalized:
            return
        self.current_block["segments"].append((normalized, self._current_style()))

    def _finalize_block(self):
        if self.current_block is None:
            return
        segments = self.current_block.get("segments") or []
        has_content = any(str(text or "").strip() for text, _ in segments)
        if has_content or self.current_block.get("tag") == "hr":
            self.blocks.append(self.current_block)
        self.current_block = None

    def handle_starttag(self, tag, attrs):
        attrs = dict(attrs)
        tag = tag.lower()

        if tag == "div" and str(attrs.get("data-page-break") or "").lower() == "true":
            self._finalize_block()
            self.blocks.append({"tag": "page_break"})
            return

        if tag in {"p", "div", "h1", "h2", "h3"}:
            self._start_block(tag, attrs)
            return

        if tag == "blockquote":
            extra_indent = self._extract_indent_points(attrs) or 36.0
            self.indent_stack.append(self._current_indent() + extra_indent)
            return

        if tag in {"ul", "ol"}:
            self.list_stack.append({"ordered": tag == "ol", "index": 0})
            return

        if tag == "li":
            self._start_block("li", attrs)
            if self.list_stack:
                current_list = self.list_stack[-1]
                if current_list["ordered"]:
                    current_list["index"] += 1
                    self._append_text(f"{current_list['index']}. ")
                else:
                    self._append_text("• ")
            return

        if tag == "br":
            self._append_text("\n")
            return

        style = self._current_style()
        pushed = False

        if tag in {"b", "strong"}:
            style["bold"] = True
            pushed = True
        elif tag in {"i", "em"}:
            style["italic"] = True
            pushed = True
        elif tag == "u":
            style["underline"] = True
            pushed = True
        elif tag in {"s", "strike"}:
            style["strike"] = True
            pushed = True
        elif tag == "font":
            if attrs.get("face"):
                style["font_family"] = attrs.get("face")
            if attrs.get("size"):
                style["font_size"] = parse_editor_font_size(
                    attrs.get("size"),
                    self.default_font_size,
                )
            pushed = True
        elif tag == "span":
            style_map = parse_inline_style(attrs.get("style"))
            font_family = style_map.get("font-family")
            if font_family:
                style["font_family"] = font_family.split(",")[0].strip().strip('"').strip("'")
            font_size = style_map.get("font-size")
            if font_size:
                match = re.search(r"(\d+(?:\.\d+)?)", font_size)
                if match:
                    style["font_size"] = int(float(match.group(1)))
            font_weight = style_map.get("font-weight", "").lower()
            if font_weight in {"bold", "600", "700", "800", "900"}:
                style["bold"] = True
            if style_map.get("font-style", "").lower() == "italic":
                style["italic"] = True
            text_decoration = style_map.get("text-decoration", "").lower()
            if "underline" in text_decoration:
                style["underline"] = True
            if "line-through" in text_decoration:
                style["strike"] = True
            pushed = True

        if pushed:
            self.style_stack.append(style)

    def handle_endtag(self, tag):
        tag = tag.lower()

        if tag in {"p", "div", "h1", "h2", "h3", "li"}:
            self._finalize_block()
            return

        if tag == "blockquote":
            if len(self.indent_stack) > 1:
                self.indent_stack.pop()
            return

        if tag in {"ul", "ol"}:
            if self.list_stack:
                self.list_stack.pop()
            return

        if tag in {"b", "strong", "i", "em", "u", "s", "strike", "font", "span"}:
            if len(self.style_stack) > 1:
                self.style_stack.pop()

    def handle_data(self, data):
        self._append_text(data)

    def handle_entityref(self, name):
        self._append_text(f"&{name};")

    def handle_charref(self, name):
        self._append_text(f"&#{name};")


def paragraph_alignment_for(value):
    value = str(value or "").strip().lower()
    if value == "center":
        return WD_ALIGN_PARAGRAPH.CENTER
    if value == "right":
        return WD_ALIGN_PARAGRAPH.RIGHT
    if value == "justify":
        return WD_ALIGN_PARAGRAPH.JUSTIFY
    return WD_ALIGN_PARAGRAPH.LEFT


def add_text_to_paragraph(paragraph, text, style, font_family, font_size):
    segments = str(text or "").split("\n")
    for index, segment in enumerate(segments):
        run = paragraph.add_run(segment)
        apply_font_to_run(
            run,
            style.get("font_family") or font_family,
            style.get("font_size") or font_size,
        )
        run.bold = bool(style.get("bold"))
        run.italic = bool(style.get("italic"))
        run.underline = bool(style.get("underline"))
        run.font.strike = bool(style.get("strike"))
        if index < len(segments) - 1:
            run.add_break()


def render_html_document_content(doc, html_content, font_family, font_size, line_spacing):
    parser = EditorHtmlParser(font_family, font_size)
    parser.feed(str(html_content or ""))
    parser.close()
    parser._finalize_block()

    written = False
    for block in parser.blocks:
        if block.get("tag") == "hr":
            paragraph = doc.add_paragraph("----------------------------------------")
            paragraph.paragraph_format.line_spacing = line_spacing
            if paragraph.runs:
                apply_font_to_run(paragraph.runs[0], font_family, font_size)
            written = True
            continue

        if block.get("tag") == "page_break":
            paragraph = doc.add_paragraph()
            paragraph.add_run().add_break(WD_BREAK.PAGE)
            written = True
            continue

        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.line_spacing = (
            block.get("line_spacing") or line_spacing
        )
        space_before_points = float(block.get("space_before") or 0.0)
        if space_before_points > 0:
            paragraph.paragraph_format.space_before = Pt(space_before_points)
        paragraph.alignment = paragraph_alignment_for(block.get("align"))
        indent_points = float(block.get("indent") or 0.0)
        if indent_points > 0:
            paragraph.paragraph_format.left_indent = Pt(indent_points)

        tag = block.get("tag")
        if tag == "h1":
            paragraph.alignment = paragraph_alignment_for(block.get("align") or "left")
        elif tag == "h2":
            paragraph.alignment = paragraph_alignment_for(block.get("align") or "left")

        segments = block.get("segments") or []
        if not segments:
            paragraph.add_run("")

        for text, style in segments:
            segment_style = dict(style)
            if tag == "h1":
                segment_style["bold"] = True
                segment_style["font_size"] = max(font_size + 8, 22)
            elif tag == "h2":
                segment_style["bold"] = True
                segment_style["font_size"] = max(font_size + 4, 18)
            elif tag == "h3":
                segment_style["bold"] = True
                segment_style["font_size"] = max(font_size + 2, 16)
            add_text_to_paragraph(
                paragraph,
                text,
                segment_style,
                font_family,
                font_size,
            )
        written = True

    return written


def replace_document_content(filepath, content, font_family, font_size, paper_size, line_spacing, margin_size, html_content=None):
    doc = DocxDocument(filepath)
    clear_document_body(doc)

    written = False
    if str(html_content or "").strip():
        written = render_html_document_content(
            doc,
            html_content,
            font_family,
            font_size,
            line_spacing,
        )

    if not written:
        paragraphs = [part.strip() for part in str(content or "").replace("\r\n", "\n").split("\n\n")]
        for paragraph_text in paragraphs:
            if not paragraph_text:
                continue
            paragraph = doc.add_paragraph(paragraph_text)
            paragraph.paragraph_format.line_spacing = line_spacing
            if paragraph.runs:
                apply_font_to_run(paragraph.runs[0], font_family, font_size)
            written = True

    if not written:
        paragraph = doc.add_paragraph("")
        paragraph.paragraph_format.line_spacing = line_spacing

    if str(html_content or "").strip():
        apply_default_font_style(doc, font_family, font_size)
        apply_document_layout_settings(
            doc,
            paper_size,
            line_spacing,
            margin_size,
            apply_line_spacing=False,
        )
    else:
        apply_font_settings(doc, font_family, font_size)
        apply_document_layout_settings(doc, paper_size, line_spacing, margin_size)
    doc.save(filepath)


def smartlegal_text_to_word_html(text):
    normalized_text = str(text or "").replace("\r\n", "\n").replace("\r", "\n")
    parts = []

    for block in re.split(r"\n\s*\n+", normalized_text):
        cleaned_block = block.strip()
        if not cleaned_block:
            continue
        block_html = escape(cleaned_block).replace("\n", "<br/>")
        parts.append(f"<p>{block_html}</p>")

    if not parts:
        parts = ["<p></p>"]
    return """<!DOCTYPE html>
<html>
  <head>
    <meta charset="UTF-8" />
    <style>
    body { font-family: "Times New Roman", serif; font-size: 14pt; line-height: 1.45; color: #111; }
    p { margin: 0 0 12px 0; }
    h1, h2, h3 { margin: 0 0 12px 0; }
  </style>
</head>
<body>
""" + "\n".join(parts) + """
</body>
</html>"""


def _cluster_column_positions(values, tolerance=22.0):
    clusters = []
    for value in sorted(float(v) for v in values):
        matched = None
        for cluster in clusters:
            if abs(cluster["center"] - value) <= tolerance:
                matched = cluster
                break
        if matched is None:
            clusters.append({"values": [value], "center": value})
        else:
            matched["values"].append(value)
            matched["center"] = sum(matched["values"]) / len(matched["values"])
    return [round(cluster["center"], 2) for cluster in clusters]


def _nearest_anchor_index(value, anchors):
    if not anchors:
        return None
    distances = [(abs(float(value) - anchor), index) for index, anchor in enumerate(anchors)]
    distance, index = min(distances, key=lambda item: item[0])
    if distance > 28.0:
        return None
    return index


def smartlegal_pdf_lines_to_word_html(pages_payload):
    blocks = []
    current_block = None

    def is_heading_line(line):
        text = str(line.get("text", "") or "").strip()
        if not text:
            return False
        letters_only = re.sub(r"[^A-Za-z]", "", text)
        is_all_caps = bool(letters_only) and letters_only == letters_only.upper()
        is_short = len(text) <= 80
        return is_short and (line.get("isBold") or is_all_caps)

    def is_list_item_line(line):
        text = str(line.get("text", "") or "").strip()
        return bool(re.match(r"^(\d+|[A-Za-z]|[IVXLC]+)[\.\)]\s+", text, flags=re.I))

    def ends_sentence(text):
        return bool(re.search(r"""[.!?;:]["')\]]?$""", text.strip()))

    def flush_block():
        nonlocal current_block
        if not current_block:
            return
        text = re.sub(r"\s+", " ", current_block["text"]).strip()
        if text:
            current_block["text"] = text
            blocks.append(current_block)
        current_block = None

    for page_index, page_data in enumerate(pages_payload):
        if page_index > 0:
            flush_block()

        for line in page_data.get("lines", []):
            text = str(line.get("text", "") or "").strip()
            if not text:
                continue

            line_kind = "heading" if is_heading_line(line) else "paragraph"
            if is_list_item_line(line):
                line_kind = "list_item"

            alignment = str(line.get("alignment", "left") or "left").lower()
            if alignment not in {"left", "center", "right", "justify"}:
                alignment = "left"

            should_start_new = current_block is None
            if not should_start_new:
                if line_kind != current_block["kind"]:
                    should_start_new = True
                elif line_kind == "heading":
                    should_start_new = True
                elif alignment != current_block["alignment"]:
                    should_start_new = True
                elif bool(line.get("isBold")) != current_block["isBold"]:
                    should_start_new = True
                elif line_kind == "list_item" and is_list_item_line(line):
                    should_start_new = True

            if should_start_new:
                flush_block()
                current_block = {
                    "text": text,
                    "kind": line_kind,
                    "alignment": alignment,
                    "isBold": bool(line.get("isBold")),
                }
                continue

            previous_text = current_block["text"].rstrip()
            if previous_text.endswith("-") and not previous_text.endswith(" -"):
                current_block["text"] = previous_text[:-1] + text
            else:
                separator = " " if not ends_sentence(previous_text) else " "
                current_block["text"] += separator + text

        flush_block()

    html_parts = []
    for block in blocks:
        text = escape(block["text"])
        if block["isBold"] and block["kind"] != "heading":
            text = f"<strong>{text}</strong>"

        style_bits = []
        if block["alignment"] != "left":
            style_bits.append(f"text-align: {block['alignment']}")
        style_attr = f' style="{"; ".join(style_bits)}"' if style_bits else ""

        if block["kind"] == "heading":
            html_parts.append(f"<h2{style_attr}><strong>{text}</strong></h2>")
        else:
            html_parts.append(f"<p{style_attr}>{text}</p>")

    if not html_parts:
        html_parts.append("<p></p>")

    return """<!DOCTYPE html>
<html>
  <head>
    <meta charset="UTF-8" />
    <style>
    body { font-family: "Times New Roman", serif; font-size: 14pt; line-height: 1.45; color: #111; }
    p { margin: 0 0 12px 0; text-align: left; }
    h1, h2, h3 { margin: 0 0 12px 0; text-align: left; }
    h2 { font-size: 18pt; }
    strong { font-weight: 700; }
  </style>
</head>
<body>
""" + "\n".join(html_parts) + """
</body>
</html>"""


def normalize_pdf_font_family(font_name):
    normalized = str(font_name or "").split("+")[-1].replace("-", " ").lower()
    if "courier" in normalized:
        return "Courier New"
    if any(name in normalized for name in ("arial", "helvetica", "calibri", "verdana", "tahoma")):
        return "Arial"
    return "Times New Roman"


def build_pdf_layout_draft(filepath):
    reader = PyPDF2.PdfReader(filepath)
    pages_payload = []
    plain_lines = []

    for page_index, page in enumerate(reader.pages):
        page_width = float(page.mediabox.width)
        page_height = float(page.mediabox.height)
        raw_spans = []

        def visitor_text(text, cm, tm, font_dict, font_size):
            value = str(text or "").replace("\r", "")
            if not value.strip():
                return

            scale_x = abs(float(cm[0] or 1.0)) or 1.0
            scale_y = abs(float(cm[3] or cm[0] or 1.0)) or 1.0
            x = float(cm[4] or 0.0) + (float(tm[4] or 0.0) * scale_x)
            baseline = float(cm[5] or 0.0) + (float(tm[5] or 0.0) * scale_y)
            font_size_value = max(8.0, float(font_size or 12.0))
            top = max(0.0, page_height - baseline - (font_size_value * 0.9))

            raw_spans.append(
                {
                    "text": value,
                    "x": x,
                    "top": top,
                    "font_size": font_size_value,
                    "font_name": str(font_dict.get("/BaseFont") if font_dict else ""),
                }
            )

        page.extract_text(visitor_text=visitor_text, Tj_sep=" ", TJ_sep=" ")
        raw_spans = [span for span in raw_spans if str(span.get("text") or "").strip()]
        raw_spans.sort(key=lambda span: (round(float(span["top"]), 1), float(span["x"])))

        grouped_lines = []
        for span in raw_spans:
            matched_line = None
            tolerance = max(3.0, float(span["font_size"]) * 0.45)
            for existing in reversed(grouped_lines[-10:]):
                if abs(float(existing["top"]) - float(span["top"])) <= tolerance:
                    matched_line = existing
                    break

            if matched_line is None:
                matched_line = {"top": float(span["top"]), "items": []}
                grouped_lines.append(matched_line)

            matched_line["items"].append(span)
            matched_line["top"] = min(float(matched_line["top"]), float(span["top"]))

        page_lines = []
        for line_index, grouped_line in enumerate(grouped_lines):
            items = sorted(grouped_line["items"], key=lambda item: float(item["x"]))
            if not items:
                continue

            text_parts = []
            previous_x = None
            for item in items:
                chunk = re.sub(r"\s+", " ", str(item["text"] or "")).strip()
                if not chunk:
                    continue

                if text_parts and previous_x is not None:
                    gap = float(item["x"]) - previous_x
                    if gap > max(6.0, float(item["font_size"]) * 0.5):
                        if not chunk.startswith((",", ".", ";", ":", ")", "]", "}")):
                            text_parts.append(" ")

                if text_parts and text_parts[-1].endswith("-"):
                    text_parts[-1] = text_parts[-1][:-1] + chunk
                else:
                    text_parts.append(chunk)

                previous_x = float(item["x"])

            line_text = "".join(text_parts).strip()
            if not line_text:
                continue

            first_item = items[0]
            font_name = str(first_item.get("font_name") or "")
            font_name_lower = font_name.lower()
            
            # Determine text alignment based on x position and line width
            line_x = float(first_item["x"])
            line_width = sum(float(item.get("font_size", 12) * 0.6 * len(str(item.get("text", "")))) for item in items)
            left_margin_threshold = 20.0  # points from left edge
            right_margin_threshold = 50.0  # points from right edge
            center_tolerance = 30.0  # tolerance for center alignment
            
            alignment = "left"  # default
            if line_x <= left_margin_threshold:
                # Check if centered
                line_center = line_x + (line_width / 2)
                page_center = page_width / 2
                if abs(line_center - page_center) <= center_tolerance:
                    alignment = "center"
                elif (line_x + line_width) >= (page_width - right_margin_threshold):
                    alignment = "justify"  # assume justified if spans most of width
                else:
                    alignment = "left"
            elif (line_x + line_width) >= (page_width - right_margin_threshold):
                alignment = "right"
            
            line_payload = {
                "id": f"page_{page_index}_line_{line_index}",
                "text": line_text,
                "x": round(float(first_item["x"]), 2),
                "y": round(float(grouped_line["top"]), 2),
                "width": round(max(120.0, page_width - float(first_item["x"]) - 24.0), 2),
                "fontSize": round(max(float(first_item["font_size"]), 8.0), 2),
                "fontFamily": normalize_pdf_font_family(font_name),
                "isBold": "bold" in font_name_lower,
                "isItalic": "italic" in font_name_lower or "oblique" in font_name_lower,
                "alignment": alignment,
            }
            page_lines.append(line_payload)
            plain_lines.append(line_text)

        pages_payload.append(
            {
                "pageIndex": page_index,
                "width": page_width,
                "height": page_height,
                "lines": page_lines,
            }
        )

    plain_text = "\n".join(plain_lines).strip()
    
    return {
        "pages": pages_payload,
        "content": plain_text,
        "html": styled_html,
    }

# -------------------------------------------------------------------
# Existing endpoint: extract_fields
# -------------------------------------------------------------------
@document_bp.route("/api/documents/upload", methods=["POST"])
def smartlegal_upload_document():
    try:
        uploaded_file = request.files.get("document")
        if uploaded_file is None or not uploaded_file.filename:
            return jsonify({"error": "No file uploaded"}), 400

        original_name = secure_filename(uploaded_file.filename)
        ext = os.path.splitext(original_name)[1].lower()
        allowed_exts = {".pdf", ".png", ".jpg", ".jpeg", ".txt", ".docx"}
        if ext not in allowed_exts:
            return jsonify({"error": "Unsupported file type"}), 400

        document_id = uuid.uuid4().hex
        stored_filename = f"{document_id}{ext}"
        stored_path = os.path.join(SMARTLEGAL_FOLDER, stored_filename)
        uploaded_file.save(stored_path)

        user = get_current_user(default=None)
        metadata = load_smartlegal_metadata()
        metadata[document_id] = {
            "id": document_id,
            "file_name": original_name,
            "stored_filename": stored_filename,
            "file_path": stored_path,
            "file_type": uploaded_file.mimetype or ext,
            "owner_username": (user.username if user else None) or "admin",
            "firm_id": user.firm_id if user else 1,
            "firm_name": (user.firm_name if user else None) or "Default Firm",
            "status": "uploaded",
            "timestamp": datetime.now().isoformat(),
        }
        save_smartlegal_metadata(metadata)

        return jsonify(
            {
                "message": "Document uploaded successfully",
                "documentId": document_id,
            }
        ), 201
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@document_bp.route("/api/documents/<document_id>/word-draft", methods=["GET"])
def smartlegal_word_draft(document_id):
    try:
        metadata = load_smartlegal_metadata()
        info = metadata.get(document_id)
        if not info:
            return jsonify({"error": "Document not found"}), 404

        user = get_current_user(default=None)
        if user and not metadata_accessible_to_user(info, user):
            return jsonify({"error": "Document not found"}), 404

        file_path = info.get("file_path")
        if not file_path or not os.path.exists(file_path):
            return jsonify({"error": "Document not found"}), 404

        ext = os.path.splitext(file_path)[1].lower()
        if ext in (".png", ".jpg", ".jpeg"):
            return jsonify({"error": "Image OCR is not configured on this backend yet"}), 400

        if ext == ".pdf":
            raw_text = extract_text_from_file(file_path, ext)
            if not raw_text or not str(raw_text).strip():
                return jsonify({"error": "Failed to extract text"}), 400

            info["extracted_text"] = raw_text
            info["status"] = "draft_ready"
            metadata[document_id] = info
            save_smartlegal_metadata(metadata)

            return jsonify(
                {
                    "message": "Word-like draft generated",
                    "content": raw_text,
                    "html": smartlegal_text_to_word_html(raw_text),
                }
            )

        raw_text = extract_text_from_file(file_path, ext)
        if not raw_text or not str(raw_text).strip():
            return jsonify({"error": "Failed to extract text"}), 400

        info["extracted_text"] = raw_text
        info["status"] = "draft_ready"
        metadata[document_id] = info
        save_smartlegal_metadata(metadata)

        return jsonify(
            {
                "message": "Word-like draft generated",
                "content": raw_text,
                "html": smartlegal_text_to_word_html(raw_text),
            }
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@document_bp.route("/api/documents/<document_id>/source", methods=["GET"])
def smartlegal_document_source(document_id):
    try:
        metadata = load_smartlegal_metadata()
        info = metadata.get(document_id)
        if not info:
            return jsonify({"error": "Document not found"}), 404

        username = get_request_username(default=None)
        user = get_current_user(default=None) if username else None
        if username and not metadata_accessible_to_user(info, user):
            return jsonify({"error": "Document not found"}), 404

        stored_filename = info.get("stored_filename")
        file_name = info.get("file_name") or stored_filename
        if not stored_filename:
            return jsonify({"error": "Document not found"}), 404

        ext = os.path.splitext(stored_filename)[1].lower()
        mimetype = "application/pdf" if ext == ".pdf" else None
        return send_from_directory(
            SMARTLEGAL_FOLDER,
            stored_filename,
            as_attachment=False,
            download_name=file_name,
            mimetype=mimetype,
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 404


@document_bp.route("/api/documents/<document_id>/generate-word-pdf", methods=["POST"])
def smartlegal_generate_word_pdf(document_id):
    try:
        metadata = load_smartlegal_metadata()
        info = metadata.get(document_id)
        if not info:
            return jsonify({"error": "Document not found"}), 404

        user = get_current_user(default=None)
        if user and not metadata_accessible_to_user(info, user):
            return jsonify({"error": "Document not found"}), 404

        payload = request.get_json(silent=True) or {}
        html_content = str(payload.get("html", "")).strip()
        if not html_content:
            return jsonify({"error": "Edited content is required"}), 400
        font_family = normalize_font_family(payload.get("fontFamily", "Times New Roman"))
        font_size = normalize_font_size(payload.get("fontSize", 14))
        line_spacing = normalize_line_spacing(payload.get("lineSpacing", 1.15))

        # Get page sizes from payload, default to A4 if not provided
        page_sizes = payload.get("pageSizes", [])
        if not page_sizes:
            paper_size = "A4"
        else:
            # Use the first page size for the document
            first_page = page_sizes[0] if page_sizes else {}
            width_mm = first_page.get("width", 595) * 25.4 / 72  # Convert points to mm
            height_mm = first_page.get("height", 842) * 25.4 / 72  # Convert points to mm
            paper_size = f"{width_mm:.1f}x{height_mm:.1f}mm"

        original_name = os.path.splitext(info.get("file_name") or "document")[0]
        safe_stem = secure_filename(original_name) or f"word_like_{document_id}"
        editable_filename = f"{safe_stem}_{document_id[:8]}.docx"
        docx_path = os.path.join(GENERATED_DOCS_FOLDER, editable_filename)

        doc = DocxDocument()
        doc.save(docx_path)

        replace_document_content(
            docx_path,
            html_to_plain_text(html_content),
            font_family=font_family,
            font_size=font_size,
            paper_size=paper_size,
            line_spacing=line_spacing,
            margin_size="Normal",
            html_content=html_content,
        )
        save_generated_html_content(editable_filename, html_content)

        pdf_path = convert_docx_to_pdf(docx_path, GENERATED_FOLDER)
        pdf_filename = os.path.basename(pdf_path)
        timestamp_iso = datetime.now().isoformat()

        generated_metadata = load_generated_metadata()
        shared_meta = {
            "owner_username": (user.username if user else None) or info.get("owner_username") or "admin",
            "firm_id": (user.firm_id if user else None) or info.get("firm_id"),
            "firm_name": (user.firm_name if user else None) or info.get("firm_name") or "Default Firm",
            "timestamp": timestamp_iso,
            "font_family": font_family,
            "font_size": font_size,
            "line_spacing": line_spacing,
            "hidden_from_dashboard": True,
            "source": "edit_document",
        }
        generated_metadata[editable_filename] = {
            **generated_metadata.get(editable_filename, {}),
            **shared_meta,
        }
        generated_metadata[pdf_filename] = {
            **generated_metadata.get(pdf_filename, {}),
            **shared_meta,
        }
        generated_metadata[os.path.splitext(editable_filename)[0] + ".editor.html"] = {
            **generated_metadata.get(os.path.splitext(editable_filename)[0] + ".editor.html", {}),
            **shared_meta,
        }
        save_generated_metadata(generated_metadata)

        return jsonify(
            {
                "message": "Word-like PDF generated successfully",
                "pdfUrl": f"{request.host_url.rstrip('/')}/download/{pdf_filename}?download=true",
                "pdfFilename": pdf_filename,
                "mode": "word-like",
            }
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@document_bp.route("/extract_fields", methods=["POST"])
def extract_fields():
    """Return backend-defined fields for a document type."""
    doc_type = normalize_document_type(request.form.get("document_type", "power_of_attorney"))
    subtype = (request.form.get("subtype") or "").strip()
    language = normalize_generation_language(request.form.get("language", "English"))
    fields = get_fields_for_document_type(doc_type, subtype=subtype)
    fields = localize_field_schema(fields, language=language)
    return jsonify(fields)


@document_bp.route("/document_subtypes", methods=["GET"])
def document_subtypes():
    """Return available subtypes for a document type (if configured)."""
    doc_type = normalize_document_type(request.args.get("document_type"))
    return jsonify(get_subtypes_for_document_type(doc_type))

# -------------------------------------------------------------------
# New endpoint: upload_reference
# -------------------------------------------------------------------
@document_bp.route("/upload_reference", methods=["POST"])
def upload_reference():
    """Upload a reference document, attach backend-defined fields, save file and metadata."""
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400
    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "Empty file"}), 400

    doc_type = normalize_document_type(request.form.get("document_type", "power_of_attorney"))
    subtype = (request.form.get("subtype") or "").strip()
    language = normalize_generation_language(request.form.get("language", "English"))
    user = get_current_user()
    ext = os.path.splitext(file.filename)[1].lower()

    # Generate unique ID and save file
    doc_id = str(uuid.uuid4())
    filename = f"{doc_id}{ext}"
    filepath = os.path.join(UPLOADS_FOLDER, filename)
    file.save(filepath)

    default_fields = get_fields_for_document_type(doc_type, subtype=subtype)
    default_fields = localize_field_schema(default_fields, language=language)
    fields = parse_serialized_fields(
        request.form.get("extracted_fields"),
        default_fields,
    )

    # Save metadata
    metadata = load_metadata()
    metadata[doc_id] = {
        "original_name": file.filename,
        "filename": filename,
        "document_type": doc_type,
        "subtype": subtype,
        "stored_fields": fields,
        "owner_username": user.username,
        "firm_id": user.firm_id,
        "firm_name": user.firm_name,
        "timestamp": datetime.now().isoformat()
    }
    save_metadata(metadata)

    return jsonify({"document_id": doc_id, "fields": fields})

# -------------------------------------------------------------------
# New endpoint: list_references
# -------------------------------------------------------------------
@document_bp.route("/list_references", methods=["GET"])
def list_references():
    """Return list of saved reference documents, optionally filtered by document_type."""
    doc_type_filter = normalize_document_type(request.args.get("document_type"))
    user = get_current_user()
    metadata = load_metadata()
    result = []
    for doc_id, info in metadata.items():
        if not metadata_accessible_to_user(info, user):
            continue
        info_doc_type = normalize_document_type(info.get("document_type"))
        if doc_type_filter and info_doc_type != doc_type_filter:
            continue
        result.append({
            "id": doc_id,
            "original_name": info["original_name"],
            "document_type": info_doc_type,
            "subtype": info.get("subtype", ""),
            "timestamp": info["timestamp"]
        })
    return jsonify(result)

# -------------------------------------------------------------------
# New endpoint: get_reference/<doc_id>
# -------------------------------------------------------------------
@document_bp.route("/get_reference/<doc_id>", methods=["GET"])
def get_reference(doc_id):
    """Return current backend-defined fields for the reference document type."""
    user = get_current_user()
    metadata = load_metadata()
    info = metadata.get(doc_id)
    if not info:
        return jsonify({"error": "Document not found"}), 404
    if not metadata_accessible_to_user(info, user):
        return jsonify({"error": "Document not found"}), 404
    doc_type = normalize_document_type(info.get("document_type", "power_of_attorney"))
    subtype = (info.get("subtype") or "").strip()
    language = normalize_generation_language(request.args.get("language", "English"))
    default_fields = get_fields_for_document_type(doc_type, subtype=subtype)
    default_fields = localize_field_schema(default_fields, language=language)
    fields = parse_serialized_fields(info.get("stored_fields"), default_fields)
    return jsonify(fields)

# -------------------------------------------------------------------
# NEW endpoint: preview reference document
# -------------------------------------------------------------------
from flask import make_response
@document_bp.route("/references/<doc_id>/view", methods=["GET"])
def view_reference(doc_id):
    username = get_request_username(default=None)
    user = get_current_user(default=None) if username else None
    metadata = load_metadata()
    info = metadata.get(doc_id)
    if not info:
        abort(404, description="Document not found")
    if username and not metadata_accessible_to_user(info, user):
        abort(404, description="Document not found")

    filename = info["filename"]
    filepath = os.path.join(UPLOADS_FOLDER, filename)

    if not os.path.exists(filepath):
        abort(404, description="File not found")

    ext = os.path.splitext(filename)[1].lower()

    # ✅ If DOCX → convert to PDF for preview
    if ext == ".docx":
        try:
            pdf_filename = filename.replace(".docx", ".pdf")
            pdf_path = os.path.join(PREVIEW_FOLDER, pdf_filename)

            # ✅ CACHE: don't reconvert if already exists
            if not os.path.exists(pdf_path):
                pdf_path = convert_docx_to_pdf(filepath, PREVIEW_FOLDER)
            response = send_from_directory(
                PREVIEW_FOLDER,
                os.path.basename(pdf_path),
                mimetype="application/pdf",
                as_attachment=False
            )
            response.headers["Access-Control-Allow-Origin"] = "*"
            response.headers["Access-Control-Allow-Headers"] = "*"
            response.headers["Access-Control-Allow-Methods"] = "GET, OPTIONS"
            return response
        except Exception as e:
            return jsonify({"error": f"Conversion failed: {str(e)}"}), 500

    # ✅ If already PDF → show directly
    elif ext == ".pdf":
        response = send_from_directory(
            UPLOADS_FOLDER,
            filename,
            mimetype="application/pdf",
            as_attachment=False
        )
        response.headers["Access-Control-Allow-Origin"] = "*"
        response.headers["Access-Control-Allow-Headers"] = "*"
        response.headers["Access-Control-Allow-Methods"] = "GET, OPTIONS"
        return response

    # ✅ TXT → show inline
    elif ext == ".txt":
        response = send_from_directory(
            UPLOADS_FOLDER,
            filename,
            mimetype="text/plain",
            as_attachment=False
        )
        response.headers["Access-Control-Allow-Origin"] = "*"
        response.headers["Access-Control-Allow-Headers"] = "*"
        response.headers["Access-Control-Allow-Methods"] = "GET, OPTIONS"
        return response

    else:
        abort(400, description="Unsupported file type")
# -------------------------------------------------------------------
# Modified endpoint: generate-document (now accepts format parameter)
# -------------------------------------------------------------------
@document_bp.route("/generate-document", methods=["POST"])
def generate_document():
    """
    Generate a final document using either:
      - a newly uploaded reference file, or
      - a previously saved reference document (by reference_id).
      - or no reference at all, in which case Gemini drafts from fields only.
    Accepts optional 'format' field: 'table' (use document-specific table template) or 'blank' (use blank template).
    """
    document_type = normalize_document_type(request.form.get("document_type"))
    fields_json = request.form.get("fields", "{}")
    # Read format parameter (default to 'table' for backward compatibility)
    format_type = request.form.get("format", "table")  # 'table' or 'blank'
    language = normalize_generation_language(request.form.get("language", "English"))
    user = get_current_user()
    font_family = normalize_font_family(request.form.get("font_family", SUPPORTED_FONT_FAMILIES[0]))
    font_size = normalize_font_size(request.form.get("font_size", DEFAULT_FONT_SIZE))

    try:
        fields = json.loads(fields_json)
    except:
        return jsonify({"error": "Invalid fields JSON"}), 400

    paper_size = normalize_paper_size(fields.get("paper_size"))
    line_spacing = normalize_line_spacing(fields.get("line_spacing"))
    margin_size = normalize_margin_size(fields.get("margin_size"))

    if document_type == "divorce_paper":
        fields = enrich_divorce_fields(fields)

    localized_fields = localize_field_values(fields, language=language)

    reference_id = request.form.get("reference_id")
    reference_text = ""
    has_reference_context = False
    field_schema = get_fields_for_document_type(document_type)
    source_reference_info = None

    # Determine source of reference text
    if reference_id:
        # Use stored reference
        metadata = load_metadata()
        info = metadata.get(reference_id)
        if not info:
            return jsonify({"error": "Reference document not found"}), 400
        if not metadata_accessible_to_user(info, user):
            return jsonify({"error": "Reference document not found"}), 400
        source_reference_info = info
        filepath = os.path.join(UPLOADS_FOLDER, info["filename"])
        ext = os.path.splitext(info["filename"])[1].lower()
        reference_text = extract_text_from_file(filepath, ext)
        if reference_text is None:
            return jsonify({"error": "Could not read stored reference file"}), 500
        has_reference_context = True
    elif "reference_file" in request.files and request.files["reference_file"].filename:
        # Use uploaded file
        file = request.files["reference_file"]
        ext = os.path.splitext(file.filename)[1].lower()
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
            file.save(tmp.name)
            tmp_path = tmp.name
        try:
            if ext == ".txt":
                with open(tmp_path, "r", encoding="utf-8") as f:
                    reference_text = f.read()
            elif ext == ".docx":
                doc = DocxDocument(tmp_path)
                reference_text = "\n".join([para.text for para in doc.paragraphs])
            elif ext == ".pdf":
                with open(tmp_path, "rb") as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    reference_text = "\n".join([page.extract_text() for page in pdf_reader.pages])
            else:
                return jsonify({"error": f"Unsupported file type: {ext}"}), 400
        except Exception as e:
            os.unlink(tmp_path)
            return jsonify({"error": f"Could not read reference file: {str(e)}"}), 400
        finally:
            os.unlink(tmp_path)
        has_reference_context = True

    # Determine which template to use based on format_type
    template_map = {
        "power_of_attorney": "power_of_attorney_template.docx",
        "gift_deed": "gift_deed_template.docx",
        "rental_agreement": "rental_agreement_template.docx",
        "partnership_deed": "partnership_deed_template.docx",
        "affidavit": "affidavit_template.docx",
        "will_and_testament": "will_testament_template.docx",
        "bail_application": "bail_application_template.docx",
        "loan_agreement": "loan_agreement_template.docx",
    }

    blank_template = "blank_template.docx"
    template_path = None

    if not has_reference_context:
        blank_path = os.path.join(TEMPLATES_FOLDER, blank_template)
        if os.path.exists(blank_path):
            template_path = blank_path
        else:
            return jsonify({"error": "Blank template not found on server"}), 500
    elif format_type == "blank":
        # Force blank template
        blank_path = os.path.join(TEMPLATES_FOLDER, blank_template)
        if os.path.exists(blank_path):
            template_path = blank_path
        else:
            return jsonify({"error": "Blank template not found on server"}), 500
    else:  # 'table' or any other value -> try to use document-specific table template
        template_filename = template_map.get(document_type)
        if template_filename:
            candidate_path = os.path.join(TEMPLATES_FOLDER, template_filename)
            if os.path.exists(candidate_path):
                template_path = candidate_path
            else:
                print(f"Table template {template_filename} not found. Falling back to blank template.")
                # Fallback to blank template
                blank_path = os.path.join(TEMPLATES_FOLDER, blank_template)
                if os.path.exists(blank_path):
                    template_path = blank_path
                else:
                    return jsonify({"error": f"No table template for '{document_type}' and blank template missing."}), 500
        else:
            # No mapping for this document type – fallback to blank
            blank_path = os.path.join(TEMPLATES_FOLDER, blank_template)
            if os.path.exists(blank_path):
                template_path = blank_path
            else:
                return jsonify({"error": f"No template mapping for '{document_type}' and blank template missing."}), 500

    # Ask Gemini to produce the final document text either from the reference
    # or from the field values alone when no reference was provided.
    try:
        if has_reference_context:
            final_document_text = fill_document_with_fields(
                reference_text,
                localized_fields,
                document_type,
                language=language,
            )
        else:
            final_document_text = generate_document_from_fields_only(
                document_type,
                localized_fields,
                field_schema=field_schema,
                language=language,
            )
            document_title = get_document_title(document_type)
            stripped_document = final_document_text.lstrip()
            normalized_start = stripped_document[:len(document_title)].upper()
            if normalized_start != document_title:
                final_document_text = f"{document_title}\n\n{stripped_document}"
    except Exception as e:
        return jsonify({"error": f"Document generation failed: {str(e)}"}), 500


    # Create unique output filename with timestamp + incremental counter
    # ----------------------------
    client_name = fields.get("principal", fields.get("client_name", "Client"))
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_filename = f"{client_name}_{document_type}_{timestamp}_generated"

    # Incremental counter to avoid duplicates
    counter = 1
    output_filename = f"{base_filename}.docx"
    while os.path.exists(os.path.join(GENERATED_DOCS_FOLDER, output_filename)):
        counter += 1
        output_filename = f"{base_filename}_{counter}.docx"

    output_path = os.path.join(GENERATED_DOCS_FOLDER, output_filename)
    try:
        # Load the server‑side template (preserves tables, formatting)
        doc = DocxDocument(template_path)
        replace_placeholders(doc, localized_fields)
        apply_font_settings(doc, font_family, font_size)
        apply_document_layout_settings(doc, paper_size, line_spacing, margin_size)
        # Split Gemini output paragraphs
        gen_paragraphs = final_document_text.split("\n\n")

        # Find first table location
        table_paragraph_index = None
        for i, block in enumerate(doc.element.body):
            if block.tag.endswith('tbl'):
                table_paragraph_index = i
                break

        # Insert text BEFORE the table
        if table_paragraph_index is not None:
            body = doc.element.body
            for text in reversed(gen_paragraphs):
                if text.strip():
                    p = doc.add_paragraph(text.strip())
                    p.paragraph_format.line_spacing = line_spacing
                    apply_font_to_run(p.runs[0], font_family, font_size)
                    body.insert(table_paragraph_index, p._element)
        else:
            # If no table exists, append normally
            for text in gen_paragraphs:
                if text.strip():
                    p = doc.add_paragraph(text.strip())
                    p.paragraph_format.line_spacing = line_spacing
                    apply_font_to_run(p.runs[0], font_family, font_size)


        doc.save(output_path)
        pdf_path = convert_docx_to_pdf(output_path, GENERATED_FOLDER)
        pdf_filename = os.path.basename(pdf_path)
        generated_metadata = load_generated_metadata()
        timestamp_iso = datetime.now().isoformat()
        generated_metadata[output_filename] = {
            "owner_username": user.username,
            "firm_id": user.firm_id,
            "firm_name": user.firm_name,
            "timestamp": timestamp_iso,
            "font_family": font_family,
            "font_size": font_size,
            "line_spacing": line_spacing,
        }
        generated_metadata[pdf_filename] = {
            "owner_username": user.username,
            "firm_id": user.firm_id,
            "firm_name": user.firm_name,
            "timestamp": timestamp_iso,
            "font_family": font_family,
            "font_size": font_size,
            "line_spacing": line_spacing,
        }
        if reference_id and source_reference_info:
            source_reference_meta = {
                "source_reference_id": reference_id,
                "source_reference_filename": source_reference_info.get("filename"),
                "source_reference_original_name": source_reference_info.get("original_name"),
            }
            generated_metadata[output_filename].update(source_reference_meta)
            generated_metadata[pdf_filename].update(source_reference_meta)
        save_generated_metadata(generated_metadata)
    except Exception as e:
        return jsonify({"error": f"Failed to create DOCX: {str(e)}"}), 500

    return jsonify({
        "docx_file": output_filename,
        "pdf_file": pdf_filename
    })

# -------------------------------------------------------------------
# Download and view endpoints (unchanged)
# -------------------------------------------------------------------
@document_bp.route("/download/<filename>", methods=["GET"])
def download_file(filename):
    try:
        username = get_request_username(default=None)
        user = get_current_user(default=None) if username else None
        ensure_generated_metadata_defaults()
        metadata = load_generated_metadata()
        info = metadata.get(filename)
        if not info:
            return jsonify({"error": "File not found"}), 404
        if username and not metadata_accessible_to_user(info, user):
            return jsonify({"error": "File not found"}), 404
        as_attachment = request.args.get('download', 'true').lower() != 'false'
        directory = get_generated_file_directory(filename)
        return send_from_directory(
            directory,
            filename,
            as_attachment=as_attachment,
            download_name=filename
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 404

@document_bp.route("/view/<filename>", methods=["GET"])
def view_pdf(filename):
    try:
        username = get_request_username(default=None)
        user = get_current_user(default=None) if username else None
        ensure_generated_metadata_defaults()
        metadata = load_generated_metadata()
        info = metadata.get(filename)
        if not info:
            return jsonify({"error": "File not found"}), 404
        if username and not metadata_accessible_to_user(info, user):
            return jsonify({"error": "File not found"}), 404
        response = send_from_directory(
            GENERATED_FOLDER,
            filename,
            mimetype="application/pdf",
            as_attachment=False
        )
        response.headers["Access-Control-Allow-Origin"] = "*"
        response.headers["Access-Control-Allow-Headers"] = "*"
        response.headers["Access-Control-Allow-Methods"] = "GET, OPTIONS"
        return response
    except Exception as e:
        return jsonify({"error": str(e)}), 404


@document_bp.route("/generated-document-content/<filename>", methods=["GET"])
def get_generated_document_content(filename):
    try:
        username = get_request_username(default=None)
        user = get_current_user(default=None) if username else None
        ensure_generated_metadata_defaults()
        metadata = load_generated_metadata()
        info = metadata.get(filename)
        if not info:
            return jsonify({"error": "File not found"}), 404
        if username and not metadata_accessible_to_user(info, user):
            return jsonify({"error": "File not found"}), 404

        ext = os.path.splitext(filename)[1].lower()
        if ext not in (".docx", ".txt", ".pdf"):
            return jsonify({"error": "Unsupported file type"}), 400

        resolved_filename = filename
        resolved_ext = ext
        filepath = os.path.join(get_generated_file_directory(filename), filename)

        matching_docx = get_matching_generated_docx_filename(filename)
        if matching_docx:
            matching_docx_path = os.path.join(GENERATED_DOCS_FOLDER, matching_docx)
            if os.path.exists(matching_docx_path):
                resolved_filename = matching_docx
                resolved_ext = ".docx"
                filepath = matching_docx_path

        if not os.path.exists(filepath):
            return jsonify({"error": "File not found"}), 404

        resolved_info = metadata.get(resolved_filename, info)
        source_reference_docx_path = get_source_reference_docx_path(resolved_info)
        content_filepath = source_reference_docx_path or filepath

        html_content = None
        if resolved_ext == ".docx":
            html_content = load_generated_html_content(resolved_filename)
            if not html_content:
                html_content = build_html_from_docx(content_filepath)

        content = extract_text_from_file(content_filepath, resolved_ext)
        if content is None:
            return jsonify({"error": "Failed to extract content"}), 500

        stored_font_family = normalize_font_family(
            resolved_info.get("font_family", SUPPORTED_FONT_FAMILIES[0])
        )
        stored_font_size = normalize_font_size(
            resolved_info.get("font_size", DEFAULT_FONT_SIZE)
        )
        stored_line_spacing = normalize_line_spacing(
            resolved_info.get("line_spacing", 1.0)
        )

        return jsonify(
            {
                "content": content,
                "html": html_content,
                "source_filename": resolved_filename,
                "font_family": stored_font_family,
                "font_size": stored_font_size,
                "line_spacing": line_spacing_label(stored_line_spacing),
            }
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@document_bp.route("/generated-document-content/<filename>", methods=["POST"])
def update_generated_document_content(filename):
    try:
        username = get_request_username(default=None)
        user = get_current_user(default=None) if username else None
        ensure_generated_metadata_defaults()
        metadata = load_generated_metadata()
        info = metadata.get(filename)
        if not info:
            return jsonify({"error": "File not found"}), 404
        if username and not metadata_accessible_to_user(info, user):
            return jsonify({"error": "File not found"}), 404

        ext = os.path.splitext(filename)[1].lower()
        editable_filename = get_matching_generated_docx_filename(filename)
        if not editable_filename or not editable_filename.lower().endswith(".docx"):
            return jsonify({"error": "Only DOCX documents can be edited"}), 400

        payload = request.get_json(silent=True) or {}
        content = str(payload.get("content", "")).strip()
        html_content = str(payload.get("html", "")).strip()
        font_family = normalize_font_family(
            payload.get("font_family", SUPPORTED_FONT_FAMILIES[0])
        )
        font_size = normalize_font_size(
            payload.get("font_size", DEFAULT_FONT_SIZE)
        )
        line_spacing = normalize_line_spacing(
            payload.get("line_spacing", 1.0)
        )
        if html_content and not content:
            content = html_to_plain_text(html_content)
        if not content:
            return jsonify({"error": "Content is required"}), 400

        docx_path = os.path.join(GENERATED_DOCS_FOLDER, editable_filename)
        if not os.path.exists(docx_path):
            return jsonify({"error": "File not found"}), 404

        editable_info = metadata.get(editable_filename, metadata.get(filename, {}))
        source_reference_docx_path = get_source_reference_docx_path(editable_info)
        edit_target_path = source_reference_docx_path or docx_path

        replace_document_content(
            edit_target_path,
            content,
            font_family=font_family,
            font_size=font_size,
            paper_size="A4",
            line_spacing=line_spacing,
            margin_size="Normal",
            html_content=html_content,
        )

        if source_reference_docx_path:
            shutil.copy2(source_reference_docx_path, docx_path)

        if html_content:
            save_generated_html_content(editable_filename, html_content)

        pdf_filename = editable_filename.replace(".docx", ".pdf")
        pdf_path = convert_docx_to_pdf(docx_path, GENERATED_FOLDER)

        timestamp_iso = datetime.now().isoformat()
        metadata[editable_filename] = {
            **metadata.get(editable_filename, {}),
            "owner_username": (user.username if user else None) or metadata.get(filename, {}).get("owner_username") or "admin",
            "firm_id": (user.firm_id if user else None) or metadata_firm_id(metadata.get(filename, {})),
            "firm_name": (user.firm_name if user else None) or metadata_firm_name(metadata.get(filename, {})),
            "timestamp": timestamp_iso,
            "font_family": font_family,
            "font_size": font_size,
            "line_spacing": line_spacing,
        }
        if source_reference_docx_path:
            metadata[editable_filename]["source_reference_filename"] = editable_info.get("source_reference_filename")
            metadata[editable_filename]["source_reference_id"] = editable_info.get("source_reference_id")
            metadata[editable_filename]["source_reference_original_name"] = editable_info.get("source_reference_original_name")
        metadata[pdf_filename] = {
            **metadata.get(pdf_filename, {}),
            "owner_username": (user.username if user else None) or metadata.get(pdf_filename, {}).get("owner_username") or "admin",
            "firm_id": (user.firm_id if user else None) or metadata_firm_id(metadata.get(pdf_filename, {})),
            "firm_name": (user.firm_name if user else None) or metadata_firm_name(metadata.get(pdf_filename, {})),
            "timestamp": timestamp_iso,
            "font_family": font_family,
            "font_size": font_size,
            "line_spacing": line_spacing,
        }
        if source_reference_docx_path:
            metadata[pdf_filename]["source_reference_filename"] = editable_info.get("source_reference_filename")
            metadata[pdf_filename]["source_reference_id"] = editable_info.get("source_reference_id")
            metadata[pdf_filename]["source_reference_original_name"] = editable_info.get("source_reference_original_name")
        html_filename = os.path.splitext(editable_filename)[0] + ".editor.html"
        metadata[html_filename] = {
            **metadata.get(html_filename, {}),
            "owner_username": (user.username if user else None) or metadata.get(html_filename, {}).get("owner_username") or "admin",
            "firm_id": (user.firm_id if user else None) or metadata_firm_id(metadata.get(html_filename, {})),
            "firm_name": (user.firm_name if user else None) or metadata_firm_name(metadata.get(html_filename, {})),
            "timestamp": timestamp_iso,
            "font_family": font_family,
            "font_size": font_size,
            "line_spacing": line_spacing,
        }
        if source_reference_docx_path:
            metadata[html_filename]["source_reference_filename"] = editable_info.get("source_reference_filename")
            metadata[html_filename]["source_reference_id"] = editable_info.get("source_reference_id")
            metadata[html_filename]["source_reference_original_name"] = editable_info.get("source_reference_original_name")
        save_generated_metadata(metadata)

        return jsonify({
            "message": "Document updated successfully",
            "docx_file": editable_filename,
            "pdf_file": os.path.basename(pdf_path),
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500






















